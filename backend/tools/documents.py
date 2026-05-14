"""
Document generation tools.
Creates Word (.docx) files for individual chapters and the final assembled book.
"""

import os
import re
from typing import List, Optional, Union, TYPE_CHECKING, Tuple
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Styling constants ──────────────────────────────────────────────────────────

FONT_NAME = "Garamond"          # Classic book font; falls back to Times New Roman if unavailable
FONT_NAME_FALLBACK = "Times New Roman"
BODY_SIZE = 12
HEADING_SIZE = 18
SUBHEADING_SIZE = 14
LINE_SPACING = 1.5              # Standard book line spacing
PAGE_WIDTH_IN  = 6.0            # inches
PAGE_HEIGHT_IN = 9.0            # inches
MARGIN_IN      = 1.0            # inches
PAGE_WIDTH  = Inches(PAGE_WIDTH_IN)
PAGE_HEIGHT = Inches(PAGE_HEIGHT_IN)
MARGIN      = Inches(MARGIN_IN)

# Usable text area inside margins
_TEXT_W = PAGE_WIDTH_IN  - 2 * MARGIN_IN   # 4.0"
_TEXT_H = PAGE_HEIGHT_IN - 2 * MARGIN_IN   # 7.0"


def _fit_image(img_path: str, max_w: float, max_h: float) -> Tuple[object, object]:
    """
    Lee las dimensiones reales de la imagen con PIL y calcula width/height (Inches)
    que caben dentro de max_w × max_h pulgadas preservando el aspect ratio.
    Si PIL no está disponible o la imagen no se puede leer, retorna (Inches(max_w), None)
    para que python-docx calcule el alto proporcional (comportamiento anterior).
    """
    try:
        from PIL import Image as _PILImage
        with _PILImage.open(img_path) as im:
            px_w, px_h = im.size
        if px_w <= 0 or px_h <= 0:
            return Inches(max_w), None
        ratio = px_w / px_h
        # Escalar para que quepa en max_w × max_h
        w = max_w
        h = w / ratio
        if h > max_h:
            h = max_h
            w = h * ratio
        return Inches(round(w, 4)), Inches(round(h, 4))
    except Exception:
        return Inches(max_w), None


def _apply_book_styles(doc: Document) -> None:
    """Apply consistent book styling to a document."""
    # Page size (6x9 — standard trade paperback)
    section = doc.sections[0]
    section.page_width = Inches(6)
    section.page_height = Inches(9)
    section.left_margin = MARGIN
    section.right_margin = MARGIN
    section.top_margin = MARGIN
    section.bottom_margin = MARGIN

    # Normal style (body text)
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_NAME
    font.size = Pt(BODY_SIZE)
    paragraph_format = style.paragraph_format
    paragraph_format.space_after = Pt(6)
    paragraph_format.line_spacing = Pt(BODY_SIZE * LINE_SPACING)
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST


def _add_chapter_heading(doc: Document, chapter_num: int, chapter_title: str) -> None:
    """Add a styled chapter heading using Word Heading styles for navigation."""
    # "CAPÍTULO N" — Heading 2 (visible en panel de navegación)
    p_num = doc.add_paragraph(style="Heading 2")
    p_num.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_num = p_num.add_run(f"CAPÍTULO {chapter_num}")
    run_num.font.name = FONT_NAME
    run_num.font.size = Pt(11)
    run_num.font.bold = False
    run_num.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Título del capítulo — Heading 1
    p_title = doc.add_paragraph(style="Heading 1")
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run(chapter_title)
    run_title.font.name = FONT_NAME
    run_title.font.size = Pt(HEADING_SIZE)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    # Separador decorativo
    p_sep = doc.add_paragraph("— ✦ —")
    p_sep.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()


def _parse_and_add_content(doc: Document, content: str, image_placements: list, subsections: list) -> None:
    """
    Parse formatted content and add to document.
    Handles [IMAGEN: desc] and [SUBTÍTULO: text] markers.
    """
    paragraphs = content.split("\n\n")
    para_count = 0

    for para_text in paragraphs:
        para_text = para_text.strip()
        if not para_text:
            continue

        # Check for image placement marker
        img_match = re.match(r'\[IMAGEN:\s*(.+?)\]', para_text, re.IGNORECASE)
        if img_match:
            desc = img_match.group(1)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"[ILUSTRACIÓN: {desc}]")
            run.font.italic = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            continue

        # Check for subsection marker
        sub_match = re.match(r'\[SUBTÍTULO:\s*(.+?)\]', para_text, re.IGNORECASE)
        if sub_match:
            subtitle_text = sub_match.group(1)
            p = doc.add_paragraph()
            run = p.add_run(subtitle_text)
            run.font.name = FONT_NAME
            run.font.size = Pt(SUBHEADING_SIZE)
            run.font.bold = True
            doc.add_paragraph()
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Inches(0.3)
        run = p.add_run(para_text.replace("\n", " "))
        run.font.name = FONT_NAME
        run.font.size = Pt(BODY_SIZE)
        para_count += 1


def _add_content_blocks(
    doc: Document,
    content_blocks: list,
    output_dir: Optional[str] = None,
    chapter_index: int = 0,
    book_title: str = "",
    genre: str = "",
    reference_image_path: str = "",
    visual_context: str = "",
    chapter_title: str = "",
    chapter_content: str = "",
    _image_warnings: Optional[list] = None,
) -> None:
    """
    Renderiza una lista de ContentBlock (TextBlock, SubtitleBlock, ImageBlock) en el doc.

    Imágenes:
    - Cada ImageBlock genera una ilustración con Ideogram y ocupa su propia página
      (salto antes + imagen centrada + salto después) para evitar cortes visuales.
    - Sin output_dir (docxes individuales del maquetador): omitidas silenciosamente.
    - visual_context:   contexto visual del libro (época, personajes físicos, estilo artístico).
    - chapter_title:    título del capítulo — ancla temática para el prompt de Ideogram.
    - chapter_content:  texto completo del capítulo — permite generar imágenes coherentes
                        con la escena real, no solo con la descripción del maquetador.
    """
    from backend.graph.utils import TextBlock, SubtitleBlock, ImageBlock
    image_counter = 0   # índice de imagen dentro del capítulo
    seen_text     = False  # True en cuanto se renderiza el primer TextBlock
    is_children = _is_children_genre(genre)

    if is_children:
        # ── Layout infantil: tabla 2 columnas (imagen | texto) ────────────────
        # Estrategia: agrupar bloques en segmentos separados por ImageBlock.
        # Cada segmento ImageBlock + los TextBlocks siguientes forma una tabla.
        # Los TextBlocks que preceden al primer ImageBlock se renderan normal.
        # La dirección de la imagen alterna por imagen para variedad visual.

        segments: list[tuple]  = []   # (image_block | None, [text_blocks])
        pending_texts: list    = []
        pending_image           = None

        for block in content_blocks:
            if isinstance(block, ImageBlock):
                if pending_image is not None:
                    # había imagen previa sin texto → guardar sola
                    segments.append((pending_image, list(pending_texts)))
                    pending_texts = []
                pending_image = block
            elif isinstance(block, (TextBlock, SubtitleBlock)):
                if pending_image is not None:
                    # texto después de imagen → forma par
                    pending_texts.append(block)
                    # Si ya tenemos suficiente texto (2+ bloques), cerrar segmento
                    if len(pending_texts) >= 2:
                        segments.append((pending_image, list(pending_texts)))
                        pending_texts = []
                        pending_image = None
                else:
                    segments.append((None, [block]))

        # Flush lo que quede
        if pending_image is not None:
            segments.append((pending_image, list(pending_texts)))
        elif pending_texts:
            for b in pending_texts:
                segments.append((None, [b]))

        # Renderizar segmentos
        for seg_idx, (img_block, txt_blocks) in enumerate(segments):
            if img_block is None:
                # Solo texto — renderizar normal
                for blk in txt_blocks:
                    if isinstance(blk, SubtitleBlock):
                        p = doc.add_paragraph(style="Heading 3")
                        run = p.add_run(blk.text)
                        run.font.name = FONT_NAME
                        run.font.size = Pt(SUBHEADING_SIZE)
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
                        doc.add_paragraph()
                    elif isinstance(blk, TextBlock):
                        p = doc.add_paragraph()
                        run = p.add_run(blk.content)
                        run.font.name = FONT_NAME
                        run.font.size = Pt(13)
                        p.paragraph_format.space_after = Pt(6)
                        p.paragraph_format.line_spacing = Pt(20)
            else:
                # Imagen centrada en página propia, texto a continuación
                if output_dir:
                    from backend.tools.cover_generator import generate_chapter_image
                    img_path, warning = generate_chapter_image(
                        description=img_block.description,
                        output_dir=output_dir,
                        chapter_index=chapter_index,
                        image_index=image_counter,
                        book_title=book_title,
                        genre=genre,
                        reference_image_path=reference_image_path,
                        visual_context=visual_context,
                        chapter_title=chapter_title,
                        chapter_content=chapter_content,
                    )
                    if warning and _image_warnings is not None:
                        _image_warnings.append(warning)
                    image_counter += 1

                    if img_path and os.path.exists(img_path):
                        try:
                            doc.add_page_break()
                            p = doc.add_paragraph()
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = p.add_run()
                            img_w, img_h = _fit_image(img_path, max_w=_TEXT_W, max_h=_TEXT_H - 0.5)
                            if img_h is not None:
                                run.add_picture(img_path, width=img_w, height=img_h)
                            else:
                                run.add_picture(img_path, width=img_w)
                            doc.add_page_break()
                        except Exception:
                            _add_image_placeholder(doc, img_block.description)
                    else:
                        _add_image_placeholder(doc, img_block.description)
                    for blk in txt_blocks:
                        if isinstance(blk, TextBlock):
                            p = doc.add_paragraph()
                            run = p.add_run(blk.content)
                            run.font.name = FONT_NAME
                            run.font.size = Pt(13)
                else:
                    _add_image_placeholder(doc, img_block.description)
        return

    # ── Layout estándar (no infantil) ────────────────────────────────────────
    for block in content_blocks:
        if isinstance(block, SubtitleBlock):
            p = doc.add_paragraph(style="Heading 3")
            run = p.add_run(block.text)
            run.font.name = FONT_NAME
            run.font.size = Pt(SUBHEADING_SIZE)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            doc.add_paragraph()

        elif isinstance(block, ImageBlock):
            if output_dir:
                from backend.tools.cover_generator import generate_chapter_image
                img_path, warning = generate_chapter_image(
                    description=block.description,
                    output_dir=output_dir,
                    chapter_index=chapter_index,
                    image_index=image_counter,
                    book_title=book_title,
                    genre=genre,
                    reference_image_path=reference_image_path,
                    visual_context=visual_context,
                    chapter_title=chapter_title,
                    chapter_content=chapter_content,
                )
                if warning and _image_warnings is not None:
                    _image_warnings.append(warning)
                image_counter += 1
                if img_path and os.path.exists(img_path):
                    try:
                        # Primera imagen sin texto previo: va pegada al título (sin salto antes).
                        # Cualquier otra imagen: página propia (salto antes).
                        if seen_text or image_counter > 1:
                            doc.add_page_break()
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run()
                        # Ajustar al área de texto preservando aspect ratio
                        img_w, img_h = _fit_image(img_path, max_w=_TEXT_W, max_h=_TEXT_H - 0.5)
                        if img_h is not None:
                            run.add_picture(img_path, width=img_w, height=img_h)
                        else:
                            run.add_picture(img_path, width=img_w)
                        doc.add_page_break()
                    except Exception:
                        _add_image_placeholder(doc, block.description)
                else:
                    _add_image_placeholder(doc, block.description)

        else:  # TextBlock
            text = block.content.replace("\n", " ").strip()
            if not text:
                continue
            seen_text = True
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Inches(0.3)
            run = p.add_run(text)
            run.font.name = FONT_NAME
            run.font.size = Pt(BODY_SIZE)


def _add_image_placeholder(doc: Document, description: str) -> None:
    """Inserta un marcador de posición para imágenes no generadas."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[ILUSTRACIÓN: {description}]")
    run.font.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_paragraph()


# ── Géneros infantiles para layout especial ────────────────────────────────────
_CHILDREN_GENRES_DOC = ["infantil", "niños", "niñas", "children", "kids",
                         "cuentos", "cuento", "álbum", "album"]


def _is_children_genre(genre: str) -> bool:
    g = genre.lower()
    return any(k in g for k in _CHILDREN_GENRES_DOC)


def _add_children_image_block(
    doc: Document,
    img_path: str,
    text_blocks: list,
    image_on_left: bool,
) -> None:
    """
    Inserta una imagen de libro infantil en tabla de 2 columnas con texto al lado.

    Diseño:
    - Imagen en columna de 1.8" (cuadrada, colores vivos)
    - Texto en columna de 2.0" fluyendo al lado
    - Alterna imagen-izquierda / imagen-derecha según image_on_left
    - Sin bordes de tabla — visualmente invisible, solo para layout
    - Fuente infantil: más grande (13pt), line-spacing más amplio
    """
    from backend.graph.utils import TextBlock, SubtitleBlock

    IMG_COL_W  = Inches(1.9)
    TEXT_COL_W = Inches(2.0)
    IMG_SIZE   = Inches(1.75)   # imagen cuadrada dentro de la celda

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"

    # Quitar todos los bordes de la tabla
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        tblBorders.append(el)
    tblPr.append(tblBorders)

    # Asignar anchos de columna
    col_widths = [IMG_COL_W, TEXT_COL_W] if image_on_left else [TEXT_COL_W, IMG_COL_W]
    for i, width in enumerate(col_widths):
        cell = table.cell(0, i)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = OxmlElement("w:tcW")
        tcW.set(qn("w:w"), str(int(width.inches * 1440)))  # twips
        tcW.set(qn("w:type"), "dxa")
        tcPr.append(tcW)

    img_col_idx  = 0 if image_on_left else 1
    text_col_idx = 1 if image_on_left else 0

    # ── Celda imagen ──────────────────────────────────────────────────────
    img_cell = table.cell(0, img_col_idx)
    img_cell.paragraphs[0].clear()
    p_img = img_cell.paragraphs[0]
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Padding interno de la celda: 4pt arriba/abajo, 6pt izq/der
    tc_img = img_cell._tc
    tcPr_img = tc_img.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", 80), ("bottom", 80), ("left", 120), ("right", 120)]:
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), str(val))
        m.set(qn("w:type"), "dxa")
        tcMar.append(m)
    tcPr_img.append(tcMar)

    # Centrar verticalmente
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), "center")
    tcPr_img.append(vAlign)

    run_img = p_img.add_run()
    try:
        run_img.add_picture(img_path, width=IMG_SIZE, height=IMG_SIZE)
    except Exception:
        run_img.add_text("[img]")

    # ── Celda texto ───────────────────────────────────────────────────────
    text_cell = table.cell(0, text_col_idx)
    text_cell.paragraphs[0].clear()

    # Padding interno: 4pt arriba/abajo, 8pt del lado de la imagen
    tc_txt = text_cell._tc
    tcPr_txt = tc_txt.get_or_add_tcPr()
    tcMar_txt = OxmlElement("w:tcMar")
    inner_pad = 160  # 8pt del lado de la imagen
    outer_pad = 80   # 4pt del lado externo
    sides = {
        "left":  inner_pad if image_on_left else outer_pad,
        "right": outer_pad if image_on_left else inner_pad,
        "top":   80,
        "bottom": 80,
    }
    for side, val in sides.items():
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), str(val))
        m.set(qn("w:type"), "dxa")
        tcMar_txt.append(m)
    tcPr_txt.append(tcMar_txt)

    first = True
    for blk in text_blocks:
        if isinstance(blk, SubtitleBlock):
            p = text_cell.add_paragraph() if not first else text_cell.paragraphs[0]
            first = False
            run = p.add_run(blk.text)
            run.font.name = FONT_NAME
            run.font.size = Pt(13)
            run.font.bold = True
        elif isinstance(blk, TextBlock):
            p = text_cell.add_paragraph() if not first else text_cell.paragraphs[0]
            first = False
            run = p.add_run(blk.content)
            run.font.name = FONT_NAME
            run.font.size = Pt(13)      # ligeramente mayor para infantil
            # Interlineado generoso para lectura infantil
            pf = p.paragraph_format
            pf.space_after = Pt(6)
            pf.line_spacing = Pt(20)

    doc.add_paragraph()  # espacio después de la tabla


def create_chapter_docx(
    output_dir: str,
    book_title: str,
    chapter_index: int,
    chapter_title: str,
    content_blocks: Optional[list] = None,
    content: Optional[str] = None,
    image_placements: Optional[list] = None,
    subsections: Optional[list] = None,
    genre: str = "",
    reference_image_path: str = "",
    visual_context: str = "",
    chapter_content: str = "",
) -> tuple:
    """
    Create a Word document for a single chapter.
    Accepts either content_blocks (list of typed blocks) or content (plain text).
    Returns (path, image_warning) where image_warning is None if all images succeeded.
    """
    os.makedirs(output_dir, exist_ok=True)

    doc = Document()
    _apply_book_styles(doc)
    _add_chapter_heading(doc, chapter_index + 1, chapter_title)

    image_warnings: list = []

    if content_blocks is not None:
        _add_content_blocks(
            doc, content_blocks,
            output_dir=output_dir,
            chapter_index=chapter_index,
            book_title=book_title,
            genre=genre,
            reference_image_path=reference_image_path,
            visual_context=visual_context,
            chapter_title=chapter_title,
            chapter_content=chapter_content,
            _image_warnings=image_warnings,
        )
    else:
        _parse_and_add_content(
            doc,
            content or "",
            image_placements or [],
            subsections or [],
        )

    # Safe filename
    safe_title = re.sub(r'[^\w\s-]', '', book_title).strip().replace(' ', '_')[:30]
    filename = f"cap{chapter_index + 1:02d}_{safe_title}.docx"
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)

    combined_warning = image_warnings[0] if image_warnings else None
    return filepath, combined_warning


def assemble_final_book(
    output_dir: str,
    title: str,
    subtitle: str,
    author_name: str,
    genre: str,
    prefacio: str,
    pagina_legal: str,
    agradecimientos: str,
    author_bio: str,
    chapters: List[dict],
    cover_image_path: Optional[str] = None,
    reference_image_path: str = "",
    visual_context: str = "",
) -> str:
    """
    Assemble the complete book into a single Word document.
    Order: Cover → Legal → Prefacio → Chapters → Sobre el Autor → Agradecimientos
    Returns path to final .docx.
    Note: cover_description (brief para diseñador) va en brief_portada.txt, nunca aquí.
    """
    os.makedirs(output_dir, exist_ok=True)
    doc = Document()
    _apply_book_styles(doc)

    # ── Cover page ────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()

    # Embed cover image if available
    if cover_image_path and os.path.exists(cover_image_path):
        try:
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_img = p_img.add_run()
            # Reservar ~3.5" para título + subtítulo + autor debajo de la imagen
            cover_w, cover_h = _fit_image(cover_image_path, max_w=_TEXT_W, max_h=_TEXT_H - 3.5)
            if cover_h is not None:
                run_img.add_picture(cover_image_path, width=cover_w, height=cover_h)
            else:
                run_img.add_picture(cover_image_path, width=cover_w)
            doc.add_paragraph()
        except Exception:
            pass  # If image fails, continue without it

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run(title.upper())
    run.font.name = FONT_NAME
    run.font.size = Pt(28)
    run.font.bold = True

    if subtitle:
        p_sub = doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_sub = p_sub.add_run(subtitle)
        run_sub.font.name = FONT_NAME
        run_sub.font.size = Pt(16)
        run_sub.font.italic = True

    doc.add_paragraph()
    p_author = doc.add_paragraph()
    p_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_a = p_author.add_run(author_name)
    run_a.font.name = FONT_NAME
    run_a.font.size = Pt(14)

    doc.add_page_break()

    # ── Legal page ─────────────────────────────────────────────────────
    # Sin heading de sección — la página legal no lleva título en libro real
    for line in pagina_legal.split("\n"):
        line = line.strip()
        # Eliminar encabezados Markdown y líneas que solo repiten el título
        if not line or re.match(r"^#{1,6}\s+", line):
            continue
        p = doc.add_paragraph()
        for segment, is_bold, is_italic in _parse_inline_markdown(line):
            if not segment:
                continue
            run = p.add_run(segment)
            run.font.name = FONT_NAME
            run.font.size = Pt(9)
            run.bold = is_bold
            run.italic = is_italic
    doc.add_page_break()

    # ── Preface ────────────────────────────────────────────────────────
    _add_section_heading(doc, "PREFACIO")
    _add_body_text(doc, prefacio, strip_title="prefacio")
    doc.add_page_break()

    # ── Chapters ───────────────────────────────────────────────────────
    from backend.graph.utils import parse_formatted_text, TextBlock
    for ch in sorted(chapters, key=lambda x: x["index"]):
        _add_chapter_heading(doc, ch["index"] + 1, ch["title"])
        content_str = ch.get("formatted_content", ch["content"])
        content_blocks = parse_formatted_text(content_str)

        # Strip duplicate title from the first TextBlock.
        # The layouter LLM sometimes repeats the chapter title at the top of the
        # formatted text (as plain text or after parse_formatted_text removes the # prefix).
        # Two cases handled:
        #   1. Title alone  → "Crecer en la Manada"          → drop block
        #   2. Title merged → "Crecer en la Manada  Aprender…" → strip prefix, keep rest
        if content_blocks and isinstance(content_blocks[0], TextBlock):
            ch_title = ch["title"].strip()
            blk_text = content_blocks[0].content.strip()
            if len(ch_title) >= 3:  # avoid false positives for very short titles
                if blk_text.lower() == ch_title.lower():
                    content_blocks.pop(0)
                elif blk_text.lower().startswith(ch_title.lower()):
                    remainder = content_blocks[0].content[len(ch_title):].lstrip()
                    if remainder:
                        content_blocks[0] = TextBlock(content=remainder)
                    else:
                        content_blocks.pop(0)

        _add_content_blocks(
            doc, content_blocks,
            output_dir=output_dir,
            chapter_index=ch["index"],
            book_title=title,
            genre=genre,
            reference_image_path=reference_image_path,
            visual_context=visual_context,
            chapter_title=ch["title"],
            chapter_content=ch.get("content", ""),
        )
        doc.add_page_break()

    # ── Sobre el Autor ─────────────────────────────────────────────────
    if author_bio and author_bio.strip():
        _add_section_heading(doc, "SOBRE EL AUTOR")
        _add_body_text(doc, author_bio, strip_title="sobre el autor")
        doc.add_page_break()

    # ── Acknowledgments ────────────────────────────────────────────────
    _add_section_heading(doc, "AGRADECIMIENTOS")
    _add_body_text(doc, agradecimientos, strip_title="agradecimientos")

    # ── Save ──────────────────────────────────────────────────────────
    safe_title = re.sub(r'[^\w\s-]', '', title).strip().replace(' ', '_')[:40]
    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    filename = f"LIBRO_FINAL_{safe_title}_{timestamp}.docx"
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)

    return filepath


def _safe_title(title: str) -> str:
    """Nombre de archivo seguro basado en el título del libro."""
    return re.sub(r'[^\w\s-]', '', title).strip().replace(' ', '_')[:40]


def assemble_final_book_markdown(
    output_dir: str,
    title: str,
    subtitle: str,
    author_name: str,
    author_bio: str,
    genre: str,
    prefacio: str,
    pagina_legal: str,
    agradecimientos: str,
    cover_image_path: Optional[str] = None,
    chapters: List[dict] = None,
    reference_image_path: str = "",
    visual_context: str = "",
) -> str:
    """Ensambla el libro final en formato Markdown (.md)."""
    from backend.graph.utils import parse_formatted_text, SubtitleBlock, ImageBlock

    lines: List[str] = []

    # Portada
    lines += [f"# {title}", ""]
    if subtitle:
        lines += [f"*{subtitle}*", ""]
    lines += [f"**{author_name}**", ""]
    if cover_image_path and os.path.exists(cover_image_path):
        lines += [f"![Portada]({os.path.basename(cover_image_path)})", ""]

    # Página legal
    if pagina_legal:
        lines += ["---", "", pagina_legal, ""]

    # Prefacio
    if prefacio:
        lines += ["---", "", "## Prefacio", "", prefacio, ""]

    # Capítulos
    for ch in (chapters or []):
        lines += ["---", "", f"## {ch['title']}", ""]
        blocks = parse_formatted_text(ch.get("formatted_content", ch.get("content", "")))
        img_idx = 0
        for block in blocks:
            if isinstance(block, SubtitleBlock):
                lines += [f"### {block.text}", ""]
            elif isinstance(block, ImageBlock):
                img_path = os.path.join(output_dir, f"img_cap{ch['index']:02d}_{img_idx:02d}.png")
                if os.path.exists(img_path):
                    lines += [f"![{block.description}]({os.path.basename(img_path)})", ""]
                else:
                    lines += [f"*[Imagen: {block.description}]*", ""]
                img_idx += 1
            else:
                if block.text.strip():
                    lines += [block.text.strip(), ""]

    # Sobre el autor
    if author_bio:
        lines += ["---", "", "## Sobre el Autor", "", author_bio, ""]

    # Agradecimientos
    if agradecimientos:
        lines += ["---", "", "## Agradecimientos", "", agradecimientos, ""]

    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    filename = f"LIBRO_FINAL_{_safe_title(title)}_{timestamp}.md"
    filepath = os.path.join(output_dir, filename)
    with open(filepath, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    return filepath


def assemble_final_book_html(
    output_dir: str,
    title: str,
    subtitle: str,
    author_name: str,
    author_bio: str,
    genre: str,
    prefacio: str,
    pagina_legal: str,
    agradecimientos: str,
    cover_image_path: Optional[str] = None,
    chapters: List[dict] = None,
    reference_image_path: str = "",
    visual_context: str = "",
) -> str:
    """Ensambla el libro final en HTML autocontenido con imágenes en base64."""
    import base64
    from backend.graph.utils import parse_formatted_text, SubtitleBlock, ImageBlock

    def _img_b64(path: str) -> str:
        with open(path, "rb") as f:
            return base64.b64encode(f.read()).decode()

    def _esc(text: str) -> str:
        return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    CSS = """
    body{font-family:Georgia,'Times New Roman',serif;max-width:720px;margin:40px auto;
         padding:0 24px;color:#1a1a1a;line-height:1.8;font-size:17px;}
    h1{font-size:2.4em;text-align:center;margin-bottom:4px;}
    h2{font-size:1.6em;margin-top:3em;border-bottom:1px solid #ccc;padding-bottom:6px;}
    h3{font-size:1.2em;margin-top:2em;}
    p{margin:0 0 1em;}
    .subtitle{text-align:center;font-style:italic;font-size:1.2em;margin-bottom:8px;}
    .author{text-align:center;font-size:1em;color:#555;margin-bottom:40px;}
    .cover img{display:block;max-width:400px;margin:20px auto;border-radius:4px;box-shadow:0 4px 16px rgba(0,0,0,.2);}
    .chapter-img{text-align:center;margin:24px 0;}
    .chapter-img img{max-width:100%;border-radius:4px;}
    .chapter-img figcaption{font-size:.85em;color:#777;margin-top:6px;font-style:italic;}
    hr{border:none;border-top:1px solid #ddd;margin:3em 0;}
    .legal{font-size:.8em;color:#888;line-height:1.5;}
    """

    body_parts: List[str] = []

    # Portada
    body_parts.append(f"<h1>{_esc(title)}</h1>")
    if subtitle:
        body_parts.append(f'<p class="subtitle">{_esc(subtitle)}</p>')
    body_parts.append(f'<p class="author">{_esc(author_name)}</p>')
    if cover_image_path and os.path.exists(cover_image_path):
        b64 = _img_b64(cover_image_path)
        body_parts.append(f'<div class="cover"><img src="data:image/png;base64,{b64}" alt="Portada"/></div>')

    # Página legal
    if pagina_legal:
        body_parts.append(f'<hr/><div class="legal"><p>{_esc(pagina_legal)}</p></div>')

    # Prefacio
    if prefacio:
        body_parts.append(f"<hr/><h2>Prefacio</h2><p>{_esc(prefacio)}</p>")

    # Capítulos
    for ch in (chapters or []):
        body_parts.append(f"<hr/><h2>{_esc(ch['title'])}</h2>")
        blocks = parse_formatted_text(ch.get("formatted_content", ch.get("content", "")))
        img_idx = 0
        for block in blocks:
            if isinstance(block, SubtitleBlock):
                body_parts.append(f"<h3>{_esc(block.text)}</h3>")
            elif isinstance(block, ImageBlock):
                img_path = os.path.join(output_dir, f"img_cap{ch['index']:02d}_{img_idx:02d}.png")
                if os.path.exists(img_path):
                    b64 = _img_b64(img_path)
                    body_parts.append(
                        f'<figure class="chapter-img">'
                        f'<img src="data:image/png;base64,{b64}" alt="{_esc(block.description)}"/>'
                        f'<figcaption>{_esc(block.description)}</figcaption></figure>'
                    )
                img_idx += 1
            else:
                if block.text.strip():
                    body_parts.append(f"<p>{_esc(block.text.strip())}</p>")

    # Sobre el autor
    if author_bio:
        body_parts.append(f"<hr/><h2>Sobre el Autor</h2><p>{_esc(author_bio)}</p>")

    # Agradecimientos
    if agradecimientos:
        body_parts.append(f"<hr/><h2>Agradecimientos</h2><p>{_esc(agradecimientos)}</p>")

    html = (
        f'<!DOCTYPE html><html lang="es"><head><meta charset="UTF-8"/>'
        f'<meta name="viewport" content="width=device-width,initial-scale=1"/>'
        f'<title>{_esc(title)}</title>'
        f'<style>{CSS}</style></head><body>'
        + "\n".join(body_parts)
        + "</body></html>"
    )

    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    filename = f"LIBRO_FINAL_{_safe_title(title)}_{timestamp}.html"
    filepath = os.path.join(output_dir, filename)
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(html)
    return filepath


def assemble_final_book_epub(
    output_dir: str,
    title: str,
    subtitle: str,
    author_name: str,
    author_bio: str,
    genre: str,
    prefacio: str,
    pagina_legal: str,
    agradecimientos: str,
    cover_image_path: Optional[str] = None,
    chapters: List[dict] = None,
    reference_image_path: str = "",
    visual_context: str = "",
) -> str:
    """Ensambla el libro final en formato EPUB usando ebooklib."""
    from ebooklib import epub
    from backend.graph.utils import parse_formatted_text, SubtitleBlock, ImageBlock

    def _esc(text: str) -> str:
        return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    CHAPTER_CSS = (
        "body{font-family:Georgia,serif;font-size:1em;line-height:1.7;margin:1em;color:#1a1a1a;}"
        "h1,h2{font-size:1.5em;margin-top:1.5em;}"
        "h3{font-size:1.1em;margin-top:1em;}"
        "p{margin:0 0 0.8em;}"
        "figure{text-align:center;margin:1.2em 0;}"
        "figure img{max-width:100%;}"
        "figcaption{font-size:0.8em;color:#777;font-style:italic;}"
    )

    book = epub.EpubBook()
    book.set_identifier(f"obra-{_safe_title(title)}")
    book.set_title(title)
    book.set_language("es")
    book.add_author(author_name)
    if subtitle:
        book.add_metadata("DC", "description", subtitle)

    css_item = epub.EpubItem(uid="style", file_name="style.css",
                              media_type="text/css", content=CHAPTER_CSS)
    book.add_item(css_item)

    # Portada
    if cover_image_path and os.path.exists(cover_image_path):
        with open(cover_image_path, "rb") as f:
            cover_bytes = f.read()
        book.set_cover("cover.png", cover_bytes)

    spine: list = ["nav"]
    toc: list = []

    def _blocks_to_xhtml(blocks, chapter_index: int) -> str:
        parts = []
        img_idx = 0
        for block in blocks:
            if isinstance(block, SubtitleBlock):
                parts.append(f"<h3>{_esc(block.text)}</h3>")
            elif isinstance(block, ImageBlock):
                img_path = os.path.join(output_dir, f"img_cap{chapter_index:02d}_{img_idx:02d}.png")
                img_uid = f"img_cap{chapter_index:02d}_{img_idx:02d}"
                if os.path.exists(img_path):
                    with open(img_path, "rb") as f:
                        img_bytes = f.read()
                    img_item = epub.EpubItem(
                        uid=img_uid,
                        file_name=f"images/{img_uid}.png",
                        media_type="image/png",
                        content=img_bytes,
                    )
                    book.add_item(img_item)
                    parts.append(
                        f'<figure><img src="images/{img_uid}.png" alt="{_esc(block.description)}"/>'
                        f'<figcaption>{_esc(block.description)}</figcaption></figure>'
                    )
                img_idx += 1
            else:
                if block.text.strip():
                    parts.append(f"<p>{_esc(block.text.strip())}</p>")
        return "\n".join(parts)

    def _make_chapter(uid: str, title_text: str, body_html: str, file_name: str) -> epub.EpubHtml:
        chap = epub.EpubHtml(title=title_text, file_name=file_name, lang="es")
        chap.content = (
            f'<html xmlns="http://www.w3.org/1999/xhtml"><head>'
            f'<link rel="stylesheet" type="text/css" href="style.css"/></head>'
            f'<body><h2>{_esc(title_text)}</h2>{body_html}</body></html>'
        )
        chap.add_item(css_item)
        return chap

    # Página legal
    if pagina_legal:
        c = _make_chapter("legal", "Información Legal",
                           f"<p>{_esc(pagina_legal)}</p>", "legal.xhtml")
        book.add_item(c); spine.append(c)

    # Prefacio
    if prefacio:
        c = _make_chapter("prefacio", "Prefacio",
                           f"<p>{_esc(prefacio)}</p>", "prefacio.xhtml")
        book.add_item(c); spine.append(c)
        toc.append(epub.Link("prefacio.xhtml", "Prefacio", "prefacio"))

    # Capítulos
    for ch in (chapters or []):
        idx = ch["index"]
        blocks = parse_formatted_text(ch.get("formatted_content", ch.get("content", "")))
        body = _blocks_to_xhtml(blocks, idx)
        uid_str = f"chapter_{idx:02d}"
        c = _make_chapter(uid_str, ch["title"], body, f"{uid_str}.xhtml")
        book.add_item(c); spine.append(c)
        toc.append(epub.Link(f"{uid_str}.xhtml", ch["title"], uid_str))

    # Sobre el autor
    if author_bio:
        c = _make_chapter("autor", "Sobre el Autor",
                           f"<p>{_esc(author_bio)}</p>", "autor.xhtml")
        book.add_item(c); spine.append(c)
        toc.append(epub.Link("autor.xhtml", "Sobre el Autor", "autor"))

    # Agradecimientos
    if agradecimientos:
        c = _make_chapter("agradecimientos", "Agradecimientos",
                           f"<p>{_esc(agradecimientos)}</p>", "agradecimientos.xhtml")
        book.add_item(c); spine.append(c)

    book.toc   = toc
    book.spine = spine
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())

    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    filename = f"LIBRO_FINAL_{_safe_title(title)}_{timestamp}.epub"
    filepath = os.path.join(output_dir, filename)
    epub.write_epub(filepath, book)
    return filepath


def _parse_inline_markdown(text: str) -> list:
    """
    Convierte markdown inline en segmentos (texto, is_bold, is_italic).
    Procesa **bold** antes que *italic* para evitar ambigüedad.
    Retorna lista de tuplas (str, bool, bool).
    """
    segments = []
    pattern = re.compile(r'\*\*(.*?)\*\*|\*(.*?)\*', re.DOTALL)
    last_end = 0
    for m in pattern.finditer(text):
        if m.start() > last_end:
            segments.append((text[last_end:m.start()], False, False))
        if m.group(1) is not None:      # **bold**
            segments.append((m.group(1), True, False))
        else:                            # *italic*
            segments.append((m.group(2), False, True))
        last_end = m.end()
    if last_end < len(text):
        segments.append((text[last_end:], False, False))
    return segments or [(text, False, False)]


def _add_section_heading(doc: Document, text: str) -> None:
    """Añade un encabezado de sección con estilo Heading 1 de Word."""
    if not text:
        return
    p = doc.add_paragraph(style="Heading 1")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    doc.add_paragraph()


def _add_body_text(doc: Document, text: str, strip_title: str = "") -> None:
    """
    Añade texto de cuerpo al documento.
    - strip_title: si el LLM repitió el título de la sección como primer párrafo, lo elimina.
    - Convierte **bold** e *italic* a runs con formato Word real.
    - Elimina encabezados Markdown (##, ###, etc.).
    """
    # Eliminar líneas de encabezado Markdown (el título ya fue añadido por _add_section_heading)
    text = re.sub(r"^#{1,6}\s+.*$", "", text, flags=re.MULTILINE)
    text = text.strip()

    # Eliminar primer párrafo si coincide con el título de sección ya impreso
    if strip_title:
        paras = [p.strip() for p in text.split("\n\n") if p.strip()]
        if paras:
            first = paras[0].strip(" \t\r\n*#").lower()
            if first == strip_title.lower() or first == strip_title.lower() + ":":
                paras = paras[1:]
        text = "\n\n".join(paras)

    for para_text in text.split("\n\n"):
        para_text = para_text.strip()
        if not para_text:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Inches(0)
        # Convertir markdown inline en runs con formato real
        for segment, is_bold, is_italic in _parse_inline_markdown(
            para_text.replace("\n", " ")
        ):
            if not segment:
                continue
            run = p.add_run(segment)
            run.font.name = FONT_NAME
            run.font.size = Pt(BODY_SIZE)
            run.bold = is_bold
            run.italic = is_italic
